import tkinter as tk
from tkinter import (
    scrolledtext, filedialog, messagebox, ttk,
    simpledialog
)
import threading
import json
import os
import re
import time
import tempfile
import math
import zipfile

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib
    matplotlib.use("Agg") # render image without opening a backend window
    import matplotlib.pyplot as plt
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

HERE        = os.path.dirname(os.path.abspath(__file__))
configjs    = os.path.join(HERE, "config.json")
charmapjs   = os.path.join(HERE, "charmap.json")

# default symbol shortcuts
# user changes are stored in charmap.json
defaultcm: dict[str, str] = {
    "+-":  "±",   "*":   "×",   "-:-": "÷",
    "=/=": "≠",   ">_":  "≥",   "<_":  "≤",
    "->":  "⇒",   "<->": "⇔",   "~~":  "≈",

    "inf":     "∞",  "aleph": "ℵ",  "alp":  "ℵ",
    "bet":     "ℶ",  "Re":    "ℝ",  "Nat":  "ℕ",
    "Com":     "ℂ",  "Int":   "ℤ",  "Rat":  "ℚ",

    "alpha":   "α",  "beta":    "β",  "gamma":   "γ",
    "delta":   "δ",  "eps":     "ε",  "epsilon": "ε",
    "zeta":    "ζ",  "eta":     "η",  "theta":   "θ",
    "tht":     "θ",  "lambda":  "λ",  "mu":      "μ",
    "nu":      "ν",  "xi":      "ξ",  "pi":      "π",
    "rho":     "ρ",  "sigma":   "σ",  "tau":     "τ",
    "phi":     "φ",  "chi":     "χ",  "psi":     "ψ",
    "omega":   "ω",  "omg":     "Ω",  "Gamma":   "Γ",
    "Delta":   "Δ",  "Theta":   "Θ",  "Lambda":  "Λ",
    "Sigma":   "Σ",  "Phi":     "Φ",  "Psi":     "Ψ",
    "Omega":   "Ω",

    "sqrt":     "√",  "cbrt":     "∛",  "integral": "∫",
    "itg":      "∫",  "iint":     "∬",  "iiint":    "∭",
    "oint":     "∮",  "sum":      "∑",  "prod":     "∏",
    "partial":  "∂",  "nabla":    "∇",

    "forall":  "∀",  "exists":  "∃",  "nexists": "∄",
    "in":      "∈",  "notin":   "∉",  "subset":  "⊂",
    "supset":  "⊃",  "union":   "∪",  "inter":   "∩",
    "empty":   "∅",  "land":    "∧",  "lor":     "∨",
    "lnot":    "¬",  "implies": "⇒",  "iff":     "⇔",

    "dgr":  "°",  "...":  "…",  "prop": "∝",
    "perp": "⊥",  "para": "∥",  "angle":"∠",
    "tri":  "△",  "deg":  "°",
}

# these keys switch into a special input mode
MODAL_CHARS = {"/", "^", "_", "x"}

# unicode superscript/subscript lookup tables
SUP = {
    "0":"⁰","1":"¹","2":"²","3":"³","4":"⁴",
    "5":"⁵","6":"⁶","7":"⁷","8":"⁸","9":"⁹",
    "+":"⁺","-":"⁻","=":"⁼","(":"⁽",")":"⁾",
    "n":"ⁿ","i":"ⁱ",
}
SUB = {
    "0":"₀","1":"₁","2":"₂","3":"₃","4":"₄",
    "5":"₅","6":"₆","7":"₇","8":"₈","9":"₉",
    "+":"₊","-":"₋","=":"₌","(":"₍",")":"₎",
    "a":"ₐ","e":"ₑ","i":"ᵢ","o":"ₒ","x":"ₓ",
}

THEMES = {
    "light": {
        "bg":         "#FAFAFA",
        "fg":         "#1A1A1A",
        "text_bg":    "#FFFFFF",
        "text_fg":    "#1A1A1A",
        "toolbar_bg": "#F0F0F0",
        "sidebar_bg": "#F5F5F5",
        "status_bg":  "#E8E8E8",
        "select_bg":  "#C8D8F0",
        "border":     "#CCCCCC",
        "accent":     "#2563EB",
        "button_bg":  "#E4E4E4",
        "button_fg":  "#1A1A1A",
    },
    "dark": {
        "bg":         "#1E1E2E",
        "fg":         "#CDD6F4",
        "text_bg":    "#1E1E2E",
        "text_fg":    "#CDD6F4",
        "toolbar_bg": "#181825",
        "sidebar_bg": "#181825",
        "status_bg":  "#11111B",
        "select_bg":  "#313244",
        "border":     "#45475A",
        "accent":     "#89B4FA",
        "button_bg":  "#313244",
        "button_fg":  "#CDD6F4",
    },
}


class MathsNotes:
    def __init__(self, root: tk.Tk):
        self.root = root # root means it's at the start of the hierarchy and all the tkinter widgets will branch off it
        self.root.title("MathsNotes") # ---- these next few lines set up the basic window stuff/info ----
        self.root.minsize(820, 540)
        self.file_path: str | None = None # set to none so the user can reset it to something else when he saves
        self.save_lock = threading.Lock() # there will be autosave, threading.lock will be used to prevent concurrent writes
        self.autosave_timer: threading.Timer | None = None # none because there is no timer at startup; when you save a file it will launch
        self.sidebar_visible = True # basically everything here is just defaults to be changed later with user input
        self.last_saved = "Never"
        self.superscript_mode = False
        self.subscript_mode = False
        self.embedded_objects: list[dict] = [] # keep track of inserted widgets so they can be saved later
        self.object_counter = 0

        self.config = self._read_json(configjs, {
            "autosave_interval": 30,
            "autosave_enabled": True,
            "font_size": 14,
            "font_family": "Times New Roman",
            "theme": "light",
        })

        self.charmap: dict[str, str] = self._read_json(charmapjs, dict(defaultcm))
        if not self.charmap:
            self.charmap = dict(defaultcm) # fall back to defaults if the file was empty

        self.max_shortcut_len = max((len(k) for k in self.charmap), default=0) # precompute so the keystroke scanner knows its lookahead window
        self.theme = THEMES.get(self.config.get("theme", "light"), THEMES["light"])

        self.build_menu()
        self.setup_toolbar()
        self.create_editor_area()
        self.make_status_bar()
        self.apply_theme()
        self.bind_shortcuts()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close) # intercepts close so it can clean up without crashing straight off
# menu

    def build_menu(self):
        bar = tk.Menu(self.root)

        fm = tk.Menu(bar, tearoff=0)
        fm.add_command(label="New",         accelerator="Ctrl+N",       command=self.new_file)
        fm.add_command(label="Open…",       accelerator="Ctrl+O",       command=self.open_file)
        fm.add_command(label="Save",        accelerator="Ctrl+S",       command=self.save_file)
        fm.add_command(label="Save As…",    accelerator="Ctrl+Shift+S", command=self.save_file_as)
        fm.add_separator()
        fm.add_command(label="Exit", command=self.on_close)
        bar.add_cascade(label="File", menu=fm)

        em = tk.Menu(bar, tearoff=0)
        em.add_command(label="Undo",       accelerator="Ctrl+Z", command=lambda: self.editor.edit_undo())
        em.add_command(label="Redo",       accelerator="Ctrl+Y", command=lambda: self.editor.edit_redo())
        em.add_separator()
        em.add_command(label="Select All", accelerator="Ctrl+A",
                       command=lambda: self.editor.tag_add("sel", "1.0", "end"))
        bar.add_cascade(label="Edit", menu=em)

        im = tk.Menu(bar, tearoff=0)
        im.add_command(label="Table…",         command=self.open_table_dialog)
        im.add_command(label="Chart…",         command=self.open_chart_dialog)
        im.add_command(label="Symbol…",        command=self.open_symbol_dialog)
        im.add_separator()
        im.add_command(label="Fraction (a/b)", command=self.open_fraction_dialog)
        bar.add_cascade(label="Insert", menu=im)

        vm = tk.Menu(bar, tearoff=0)
        vm.add_command(label="Toggle Sidebar", accelerator="Ctrl+B", command=self.sidebar_on)
        vm.add_command(label="Light Theme",    command=lambda: self.set_theme("light"))
        vm.add_command(label="Dark Theme",     command=lambda: self.set_theme("dark"))
        bar.add_cascade(label="View", menu=vm)

        sm = tk.Menu(bar, tearoff=0)
        sm.add_command(label="Settings…",           command=self.open_settings_dialog)
        sm.add_command(label="Edit Character Map…", command=self.open_charmap_dialog)
        bar.add_cascade(label="Settings", menu=sm)

        hm = tk.Menu(bar, tearoff=0)
        hm.add_command(label="Help / Manual", accelerator="F1", command=self.show_help)
        hm.add_command(label="About",                            command=self.show_abt)
        bar.add_cascade(label="Help", menu=hm)

        self.root.config(menu=bar)
# toolbar

    def setup_toolbar(self):
        self.toolbar = tk.Frame(self.root, height=36, relief="flat", bd=0)
        self.toolbar.pack(side="top", fill="x")

        tk.Label(self.toolbar, text="Font:").pack(side="left", padx=(8, 2))
        self.font_family_var = tk.StringVar(value=self.config.get("font_family", "Times New Roman"))
        font_family_cb = ttk.Combobox(self.toolbar, textvariable=self.font_family_var,
                             values=["Times New Roman", "Georgia", "Palatino Linotype",
                                     "Courier New", "Arial", "Cambria Math", "DejaVu Serif"],
                             width=18, state="readonly")
        font_family_cb.pack(side="left", padx=2)
        font_family_cb.bind("<<ComboboxSelected>>", lambda e: self.apply_font_globally())

        tk.Label(self.toolbar, text="Size:").pack(side="left", padx=(8, 2))
        self.font_size_var = tk.IntVar(value=self.config.get("font_size", 14))
        font_size_spin = tk.Spinbox(self.toolbar, from_=8, to=72, width=4,
                           textvariable=self.font_size_var, command=self.apply_font_to_selection)
        font_size_spin.pack(side="left", padx=2)
        font_size_spin.bind("<Return>", lambda e: self.apply_font_to_selection())

        ttk.Separator(self.toolbar, orient="vertical").pack(side="left", fill="y", padx=6, pady=4)

        for label, fn in [("B", self.toggle_bold), ("I", self.toggle_italic)]:
            tk.Button(self.toolbar, text=label, width=2, relief="flat", command=fn,
                      font=("Arial", 10, "bold" if label == "B" else "italic")
                      ).pack(side="left", padx=1)

        ttk.Separator(self.toolbar, orient="vertical").pack(side="left", fill="y", padx=6, pady=4)

        tk.Label(self.toolbar, text="Notation:").pack(side="left", padx=(4, 2))
        self.notation_var = tk.StringVar(value="Insert symbol…")
        self.notation_cb = ttk.Combobox(self.toolbar, textvariable=self.notation_var, width=24, state="normal")
        self.reload_notation_combobox()
        self.notation_cb.pack(side="left", padx=2)
        self.notation_cb.bind("<<ComboboxSelected>>", self.insert_from_combobox)
        self.notation_cb.bind("<Return>",              self.insert_from_combobox)
        self.notation_cb.bind("<KeyRelease>",          self.filter_notation_combobox)

        ttk.Separator(self.toolbar, orient="vertical").pack(side="left", fill="y", padx=6, pady=4)
        tk.Button(self.toolbar, text="⊞ Table", relief="flat", command=self.open_table_dialog).pack(side="left", padx=2)
        tk.Button(self.toolbar, text="📈 Chart", relief="flat", command=self.open_chart_dialog).pack(side="left", padx=2)
        tk.Button(self.toolbar, text="≡ Ref",   relief="flat", command=self.sidebar_on).pack(side="right", padx=8)

    def reload_notation_combobox(self, query=""):
        # rebuild the dropdown list, optionally filtered by a search string
        query = query.lower().strip()
        self.notation_cb["values"] = [
            f"{k}  →  {v}" for k, v in self.charmap.items()
            if not query or query in k.lower() or query in v.lower()
        ]

    def filter_notation_combobox(self, ev=None):
        # live filter as the user types, but don't interfere with navigation keys
        if ev and ev.keysym in ("Up", "Down", "Return", "Escape"):
            return
        typed = self.notation_cb.get()
        if " →" not in typed: # only filter when the user is typing, not after a selection
            self.reload_notation_combobox(typed)
# body

    def create_editor_area(self):
        self.paned = tk.PanedWindow(self.root, orient="horizontal", sashrelief="flat", sashwidth=4)
        self.paned.pack(fill="both", expand=True)

        editor_frame = tk.Frame(self.paned)
        self.editor = scrolledtext.ScrolledText(
            editor_frame, wrap="word", undo=True, maxundo=-1,
            font=(self.config.get("font_family", "Times New Roman"), self.config.get("font_size", 14)),
            relief="flat", bd=0, padx=12, pady=12, insertwidth=2,
        )
        self.editor.pack(fill="both", expand=True)
        self.paned.add(editor_frame, stretch="always", minsize=400)

        self.sidebar_frame = tk.Frame(self.paned, width=240)
        tk.Label(self.sidebar_frame, text="Character Map", font=("Arial", 10, "bold"),
                 anchor="w").pack(fill="x", padx=8, pady=(8, 4))

        self.sidebar_query_var = tk.StringVar()
        sidebar_search = tk.Entry(self.sidebar_frame, textvariable=self.sidebar_query_var)
        sidebar_search.pack(fill="x", padx=8, pady=(0, 4))
        sidebar_search.bind("<KeyRelease>", lambda e: self.refill_sidebar(self.sidebar_query_var.get()))

        sidebar_scrollbar = tk.Scrollbar(self.sidebar_frame)
        sidebar_scrollbar.pack(side="right", fill="y")
        self.sidebar_listbox = tk.Listbox(self.sidebar_frame, yscrollcommand=sidebar_scrollbar.set,
                                 font=("Courier New", 10), relief="flat", bd=0,
                                 activestyle="none", selectmode="browse")
        self.sidebar_listbox.pack(fill="both", expand=True, padx=4)
        sidebar_scrollbar.config(command=self.sidebar_listbox.yview)
        self.refill_sidebar()
        self.sidebar_listbox.bind("<Double-Button-1>", self.insert_from_sidebar)
        self.paned.add(self.sidebar_frame, stretch="never", minsize=220)

        self.editor.bind("<KeyRelease>",     self.on_key_release)
        self.editor.bind("<KeyPress-slash>", self.on_slash_pressed) # it intercepts whenever / is pressed so it can decide whether to keep as slash or convert to 
                                                                    # fraction (see on_slash_pressed linked command)

    def refill_sidebar(self, query=""):
        self.sidebar_listbox.delete(0, "end")
        self.sidebar_listbox.insert("end", "  Shortcut       Symbol")
        self.sidebar_listbox.insert("end", "  " + "─"*24)
        query = query.lower().strip()
        for shortcut, symbol in self.charmap.items():
            if not query or query in shortcut.lower() or query in symbol.lower():
                self.sidebar_listbox.insert("end", f"  {shortcut:<14} {symbol}")
        if not query: # only show modal trigger section when not filtering
            self.sidebar_listbox.insert("end", "")
            self.sidebar_listbox.insert("end", "  Modal triggers")
            self.sidebar_listbox.insert("end", "  " + "─"*24)
            self.sidebar_listbox.insert("end", "  /              fraction")
            self.sidebar_listbox.insert("end", "  ^              superscript")
            self.sidebar_listbox.insert("end", "  _              subscript")
            self.sidebar_listbox.insert("end", "  x              × or 𝑥")

    def make_status_bar(self):
        self.status_bar = tk.Frame(self.root, height=22, relief="flat")
        self.status_bar.pack(side="bottom", fill="x")
        self.status_label = tk.Label(self.status_bar, text="Ready", anchor="w", font=("Arial", 9), padx=8)
        self.status_label.pack(side="left")
        self.save_status_label = tk.Label(self.status_bar, text="Auto-save: —", anchor="e", font=("Arial", 9), padx=8)
        self.save_status_label.pack(side="right")
# theme

    def apply_theme(self):
        t = self.theme
        self.root.configure(bg=t["bg"])
        self.toolbar.configure(bg=t["toolbar_bg"])
        self.status_bar.configure(bg=t["status_bg"])
        self.status_label.configure(bg=t["status_bg"], fg=t["fg"])
        self.save_status_label.configure(bg=t["status_bg"], fg=t["fg"])
        self.sidebar_frame.configure(bg=t["sidebar_bg"])
        self.editor.configure(bg=t["text_bg"], fg=t["text_fg"],
                          insertbackground=t["fg"], selectbackground=t["select_bg"])
        self.sidebar_listbox.configure(bg=t["sidebar_bg"], fg=t["fg"], selectbackground=t["select_bg"])
        for widget in self.toolbar.winfo_children():
            try:
                widget.configure(bg=t["toolbar_bg"], fg=t["fg"])
            except Exception:
                pass # ttk widgets silently reject the call, that's fine

    def set_theme(self, name: str):
        self.theme = THEMES.get(name, THEMES["light"])
        self.config["theme"] = name
        self.apply_theme()
        self.save_config()
# keys

    def bind_shortcuts(self):
        self.root.bind("<Control-n>", lambda e: self.new_file())
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-s>", lambda e: self.save_file())
        self.root.bind("<Control-S>", lambda e: self.save_file_as())
        self.root.bind("<Control-b>", lambda e: self.sidebar_on())
        self.root.bind("<F1>",        lambda e: self.show_help())
# realtime subs

    def get_char_before_cursor(self, n=1) -> str:
        # helper to safely peek at the n-th character before the cursor
        try:
            return self.editor.get(f"insert - {n} chars", f"insert - {n-1} chars")
        except tk.TclError:
            return ""

    def on_key_release(self, ev):
        # try to evaluate a simple line when Enter is pressed
        if ev.keysym == "Return":
            self.superscript_mode = self.subscript_mode = False # pressing enter always resets modal modes
            line_start = self.editor.index("insert - 1 lines linestart")
            line_end   = self.editor.index("insert - 1 lines lineend")
            result = try_evaluate_equation(self.editor.get(line_start, line_end).strip())
            if result is not None:
                self.editor.insert("insert - 1 chars", f" {result}")

        if ev.keysym in ("Shift_L","Shift_R","Control_L","Control_R",
                          "Alt_L","Alt_R","Left","Right","Up","Down",
                          "Home","End","Escape","Tab"):
            return # navigation and modifier keys can't form shortcuts so bail early

        ch = self.get_char_before_cursor(1)

        if self.superscript_mode and ch and ch != "^":
            if ch in SUP:
                self.editor.delete("insert - 1 chars", "insert")
                self.editor.insert("insert", SUP[ch]) # swap raw char for unicode superscript
            else:
                self.superscript_mode = False
                self.update_status("super: off") # unmappable char exits the mode
            return

        if self.subscript_mode and ch and ch != "_":
            if ch in SUB:
                self.editor.delete("insert - 1 chars", "insert")
                self.editor.insert("insert", SUB[ch]) # swap raw char for unicode subscript
            else:
                self.subscript_mode = False
                self.update_status("sub: off")
            return

        if ch == "^":
            self.editor.delete("insert - 1 chars", "insert") # eat the ^ so it doesn't appear in the doc
            self.superscript_mode = True
            self.update_status("super: on  (non-digit exits)")
            return

        if ch == "_":
            two_chars = self.editor.get("insert - 2 chars", "insert")
            if two_chars not in (">_", "<_"): # >_ and <_ are charmap shortcuts, don't steal the _
                self.editor.delete("insert - 1 chars", "insert")
                self.subscript_mode = True
                self.update_status("sub: on")
                return

        # check for symbol shortcuts near the cursor, scan longest-1st so "epsilon" beats "eps"
        cursor_pos = self.editor.index("insert")
        for n in range(self.max_shortcut_len, 0, -1):
            try:
                start = self.editor.index(f"insert - {n} chars")
            except tk.TclError:
                continue
            seq = self.editor.get(start, cursor_pos)
            if seq in self.charmap:
                self.editor.delete(start, cursor_pos)
                self.editor.insert(start, self.charmap[seq])
                return

        self.update_status()

    def on_slash_pressed(self, ev):
        # if the user types / after a number or bracket, make a fraction
        # this is a KeyPress binding so it fires before the / appears in the widget
        try:
            cur  = self.editor.index("insert")
            line = self.editor.get(f"{cur} linestart", cur)
        except tk.TclError:
            return None

        if not line:
            return None

        tail = line[-1]
        if not (tail.isdigit() or tail in (")", "]", "}")):
            return None # only trigger after a digit or closing bracket

        if tail in (")", "]", "}"): # catches whenever / should be a fraction; here it looks for brackets as a sign of a maths clause
            opener = {")":"(", "]":"[", "}":"{"}[tail]
            depth, i = 0, len(line)-1
            hit = -1 # hit = index where matching bracket is found, so it starts from the end
            while i >= 0: # walk backwards to find the matching opener
                if line[i] == tail:    depth += 1 # increases depth by 1
                elif line[i] == opener:
                    depth -= 1
                    if depth == 0: hit = i; break # marks hit to exit
                i -= 1
            if hit < 0:
                return None # hit stays the same (-1 which is impossible) means it messed up
            numerator_text = line[hit:]
        else:
            bounds = set("+-*×÷=<>,;:|&^~ \t([{")
            i = len(line)-1
            while i >= 0 and line[i] not in bounds:
                i -= 1
            numerator_text = line[i+1:]

        if not numerator_text:
            return None

        num_start = self.editor.index(f"insert - {len(numerator_text)} chars")
        self.editor.delete(num_start, cur) # remove the numerator text before replacing with the widget
        self.fractionmaker(num_start, numerator_text)
        return "break" # suppress the / from being inserted
# fraction widget

    def fractionmaker(self, at: str, num="", den=""):
        t   = self.theme
        fam = self.config.get("font_family", "Times New Roman")
        sz  = max(8, self.config.get("font_size", 14) - 2) # slightly smaller than body text so it sits inline comfortably

        box = tk.Frame(self.editor, background=t["text_bg"], bd=0, highlightthickness=0)
        if not hasattr(self, "fraction_frames"):
            self.fraction_frames: list[tk.Frame] = []
        self.fraction_frames.append(box)

        num_var = tk.StringVar(value=num)
        den_var = tk.StringVar(value=den)

        def calc_width():
            return max(2, len(num_var.get()), len(den_var.get())) + 1

        initial_width = calc_width()
        entry_kwargs = dict(font=(fam, sz), relief="flat", bd=0, highlightthickness=0,
                  background=t["text_bg"], foreground=t["text_fg"],
                  insertbackground=t["fg"], justify="center", width=initial_width)
        num_entry = tk.Entry(box, textvariable=num_var, **entry_kwargs)
        num_entry.grid(row=0, column=0, sticky="ew", padx=2)
        tk.Frame(box, height=1, background=t["fg"]).grid(row=1, column=0, sticky="ew", padx=1, pady=1) # fraction bar
        den_entry = tk.Entry(box, textvariable=den_var, **entry_kwargs)
        den_entry.grid(row=2, column=0, sticky="ew", padx=2)

        def on_content_change(*_):
            # resize both entries so they always match the wider of the two
            new_width = calc_width()
            num_entry.configure(width=new_width)
            den_entry.configure(width=new_width)
            box.update_idletasks()
        num_var.trace_add("write", on_content_change)
        den_var.trace_add("write", on_content_change)

        def return_to_editor(ev=None):
            self.editor.focus_set()
            return "break"
        for entry in (num_entry, den_entry):
            entry.bind("<Return>", return_to_editor)
            entry.bind("<Tab>",    return_to_editor)
            entry.bind("<Escape>", return_to_editor)

        self.editor.window_create(at, window=box, align="center") # embed the frame inline with the text

        mark_name = f"obj_{self.object_counter}"
        self.object_counter += 1
        self.editor.mark_set(mark_name, at)
        self.editor.mark_gravity(mark_name, "left") # left gravity keeps the mark to the left of the widget as text grows
        self.embedded_objects.append({
            "type": "fraction", "mark": mark_name, "widget": box,
            "getter": lambda n=num_var, d=den_var: {"numerator": n.get(), "denominator": d.get()},
        })
        den_entry.focus_set() # move focus to denominator so the user can type it immediately

    def open_fraction_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Insert Fraction"); dlg.transient(self.root)
        dlg.grab_set(); dlg.resizable(False, False)
        tk.Label(dlg, text="Numerator:").grid(  row=0, column=0, padx=12, pady=(10,4), sticky="e")
        num_var = tk.StringVar()
        num_entry = tk.Entry(dlg, textvariable=num_var, width=20)
        num_entry.grid(row=0, column=1, padx=8, pady=(10,4))
        num_entry.focus_set()
        tk.Label(dlg, text="Denominator:").grid(row=1, column=0, padx=12, pady=4, sticky="e")
        den_var = tk.StringVar()
        den_entry = tk.Entry(dlg, textvariable=den_var, width=20)
        den_entry.grid(row=1, column=1, padx=8, pady=4)
        def do_insert(ev=None):
            n = num_var.get().strip()
            d = den_var.get().strip()
            dlg.destroy()
            self.fractionmaker("insert", n, d)
        num_entry.bind("<Return>", lambda e: den_entry.focus_set())
        den_entry.bind("<Return>", do_insert)
        tk.Button(dlg, text="Insert", width=10, command=do_insert).grid(row=2, column=0, columnspan=2, pady=10)
# fonts

    def make_font_tag_name(self, bold=False, italic=False):
        # build a unique tag name encoding the current font family, size, and style
        parts = ["fontstyle", self.font_family_var.get().replace(" ", "_"), str(self.font_size_var.get())]
        if bold:   parts.append("bold")
        if italic: parts.append("italic")
        return "_".join(parts)

    def configure_font_tag(self, tag, bold=False, italic=False):
        font_parts = [self.font_family_var.get(), self.font_size_var.get()]
        if bold:   font_parts.append("bold")
        if italic: font_parts.append("italic")
        self.editor.tag_configure(tag, font=tuple(font_parts))

    def strip_style_tags(self, start, end):
        # remove all font-related tags from a range so we can reapply cleanly
        for tag in self.editor.tag_names():
            if tag.startswith("font_") or tag.startswith("fontstyle_") or tag in ("bold","italic"):
                self.editor.tag_remove(tag, start, end)

    def rebuild_style_tags(self):
        # after a font change, reconstruct all bold/italic compound tags with the new family/size
        end = self.editor.index("end-1c")
        if self.editor.compare(end, "<=", "1.0"):
            return
        bold_ranges   = list(self.editor.tag_ranges("bold"))
        italic_ranges = list(self.editor.tag_ranges("italic"))
        for tag in list(self.editor.tag_names()):
            if tag.startswith("fontstyle_"):
                self.editor.tag_delete(tag) # wipe old compound tags before rebuilding
        all_points = sorted({"1.0", end} | {str(x) for x in bold_ranges + italic_ranges},
                     key=lambda s: tuple(map(int, s.split("."))))
        for a, b in zip(all_points, all_points[1:]):
            if self.editor.compare(a, ">=", b): continue
            is_bold   = "bold"   in self.editor.tag_names(a)
            is_italic = "italic" in self.editor.tag_names(a)
            if not (is_bold or is_italic): continue
            tag = self.make_font_tag_name(bold=is_bold, italic=is_italic)
            self.configure_font_tag(tag, bold=is_bold, italic=is_italic)
            self.editor.tag_add(tag, a, b)
            self.editor.tag_raise(tag)

    def refresh_embedded_fonts(self):
        # push the current font choice into every live fraction and table widget
        fam  = self.font_family_var.get()
        sz   = self.font_size_var.get()
        small = max(8, sz-2) # fractions use a slightly smaller size so they sit comfortably inline
        for obj in getattr(self, "embedded_objects", []):
            widget = obj.get("widget")
            if widget is None: continue
            try:
                if not widget.winfo_exists(): continue
            except tk.TclError:
                continue
            if obj.get("type") == "fraction":
                for child in widget.winfo_children():
                    if isinstance(child, tk.Entry):
                        child.configure(font=(fam, small))
            elif obj.get("type") == "table":
                has_header = obj.get("has_header", False)
                for child in widget.winfo_children():
                    grid_info = child.grid_info()
                    row_index = int(grid_info.get("row", 0)) if grid_info else 0
                    if isinstance(child, tk.Entry):
                        child.configure(font=(fam, sz, "bold") if (has_header and row_index == 0) else (fam, sz))
        self.editor.update_idletasks()

    def apply_font_globally(self):
        # apply current font family and size to the whole document
        self.editor.configure(font=(self.font_family_var.get(), self.font_size_var.get()))
        self.editor.tag_configure("bold")
        self.editor.tag_configure("italic")
        self.rebuild_style_tags()
        self.refresh_embedded_fonts()
        self.config["font_family"] = self.font_family_var.get()
        self.config["font_size"]   = self.font_size_var.get()
        self.save_config()

    def apply_font_to_selection(self):
        fam = self.font_family_var.get()
        sz  = self.font_size_var.get()
        try:
            sel_start = self.editor.index("sel.first")
            sel_end   = self.editor.index("sel.last")
        except tk.TclError:
            self.apply_font_globally() # nothing selected, apply everywhere
            return
        is_bold   = "bold"   in self.editor.tag_names(sel_start)
        is_italic = "italic" in self.editor.tag_names(sel_start)
        self.strip_style_tags(sel_start, sel_end)
        if is_bold:   self.editor.tag_add("bold",   sel_start, sel_end)
        if is_italic: self.editor.tag_add("italic", sel_start, sel_end)
        tag = self.make_font_tag_name(bold=is_bold, italic=is_italic)
        self.configure_font_tag(tag, bold=is_bold, italic=is_italic)
        self.editor.tag_add(tag, sel_start, sel_end)
        self.editor.tag_raise(tag)
        self.config["font_family"] = fam
        self.config["font_size"]   = sz
        self.refresh_embedded_fonts()
        self.save_config()

    def toggle_bold(self):
        try:
            sel_start = self.editor.index("sel.first")
            sel_end   = self.editor.index("sel.last")
        except tk.TclError:
            return # nothing selected, do nothing
        is_bold   = "bold"   in self.editor.tag_names(sel_start)
        is_italic = "italic" in self.editor.tag_names(sel_start)
        self.strip_style_tags(sel_start, sel_end)
        is_bold = not is_bold # flip the bold state
        if is_bold:   self.editor.tag_add("bold",   sel_start, sel_end)
        else:         self.editor.tag_remove("bold", sel_start, sel_end)
        if is_italic: self.editor.tag_add("italic", sel_start, sel_end)
        if is_bold or is_italic:
            tag = self.make_font_tag_name(bold=is_bold, italic=is_italic)
            self.configure_font_tag(tag, bold=is_bold, italic=is_italic)
            self.editor.tag_add(tag, sel_start, sel_end)
            self.editor.tag_raise(tag)

    def toggle_italic(self):
        try:
            sel_start = self.editor.index("sel.first")
            sel_end   = self.editor.index("sel.last")
        except tk.TclError:
            return # nothing selected, do nothing
        is_bold   = "bold"   in self.editor.tag_names(sel_start)
        is_italic = "italic" in self.editor.tag_names(sel_start)
        self.strip_style_tags(sel_start, sel_end)
        is_italic = not is_italic # flip the italic state
        if is_italic: self.editor.tag_add("italic", sel_start, sel_end)
        else:         self.editor.tag_remove("italic", sel_start, sel_end)
        if is_bold:   self.editor.tag_add("bold",   sel_start, sel_end)
        if is_bold or is_italic:
            tag = self.make_font_tag_name(bold=is_bold, italic=is_italic)
            self.configure_font_tag(tag, bold=is_bold, italic=is_italic)
            self.editor.tag_add(tag, sel_start, sel_end)
            self.editor.tag_raise(tag)
# file ops

    def new_file(self):
        if not self.check_unsaved(): return
        self.editor.delete("1.0", "end")
        self.embedded_objects.clear()
        self.object_counter = 0
        self.file_path = None
        self.root.title("MathsNotes — Untitled")
        self.update_status("New file")
        self.stop_autosave()

    SIDECAR_PART = "customXml/mathsnotes.json" # stored inside the docx zip, invisible to Word
    SIDECAR_VER  = 2

    def open_file(self):
        if not self.check_unsaved(): return
        if not HAS_DOCX:
            messagebox.showerror("Missing Dependency",
                "MathsNotes needs python-docx for .docx files.\n\nInstall with:\n  pip install python-docx")
            return
        path = filedialog.askopenfilename(title="Open Document",
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
        if not path: return
        try:
            self.load_file(path)
            self.file_path = path
            self.root.title(f"MathsNotes — {os.path.basename(path)}")
            self.update_status(f"Opened: {os.path.basename(path)}")
            self.start_autosave()
        except Exception as e:
            messagebox.showerror("Open Error", f"Could not open file:\n{e}")

    def load_file(self, path: str):
        self.editor.delete("1.0", "end")
        self.embedded_objects.clear()
        self.object_counter = 0

        sidecar = None
        try:
            with zipfile.ZipFile(path, "r") as zf:
                if self.SIDECAR_PART in zf.namelist():
                    sidecar = json.loads(zf.read(self.SIDECAR_PART).decode("utf-8"))
        except (zipfile.BadZipFile, KeyError, json.JSONDecodeError):
            pass # if the sidecar is missing or broken, open as plain text

        # build a lookup from body child index → object descriptor for fast reconstruction
        objects_by_body_index: dict[int, dict] = {}
        if sidecar and isinstance(sidecar.get("objects"), list):
            for obj in sidecar["objects"]:
                if isinstance(obj.get("body_index"), int):
                    objects_by_body_index[obj["body_index"]] = obj

        doc   = Document(path)
        first = True
        for idx, child in enumerate(doc.element.body.iterchildren()):
            tag      = child.tag.split("}")[-1]
            sidecar_obj = objects_by_body_index.get(idx)

            if sidecar_obj and sidecar_obj.get("type") == "fraction":
                self.fractionmaker("insert", sidecar_obj.get("numerator",""), sidecar_obj.get("denominator",""))
                self.editor.focus_set(); first = False; continue

            if sidecar_obj and sidecar_obj.get("type") == "table":
                self.make_table_widget(sidecar_obj.get("data",[]), sidecar_obj.get("has_header", False))
                first = False; continue

            if tag == "p":
                text = "".join(t.text or "" for t in child.iter(qn("w:t")))
                if not first: self.editor.insert("insert", "\n")
                self.editor.insert("insert", text)
                first = False
            elif tag == "tbl":
                rows = []
                for tr in child.iter(qn("w:tr")):
                    cells = ["".join(t.text or "" for t in tc.iter(qn("w:t")))
                             for tc in tr.iter(qn("w:tc"))]
                    if cells: rows.append(cells)
                if rows: self.make_table_widget(rows, False)
                first = False

    def save_file(self):
        with self.save_lock: # lock prevents autosave and manual save from writing simultaneously
            if self.file_path is None:
                path = filedialog.asksaveasfilename(title="Save As",
                        defaultextension=".docx",
                        filetypes=[("Word documents", "*.docx")])
                if not path: return
                self.file_path = path
                self.start_autosave()
            self.write_file(self.file_path)

    def save_file_as(self):
        path = filedialog.asksaveasfilename(title="Save As",
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")])
        if not path: return
        self.file_path = path
        with self.save_lock: self.write_file(path)
        self.start_autosave()

    def get_document_segments(self) -> list[dict]:
        # walk the editor content and split it into alternating text and widget segments
        live_objects: list[tuple[str, dict]] = []
        for obj in self.embedded_objects:
            try:
                if not obj["widget"].winfo_exists(): continue
                live_objects.append((self.editor.index(obj["mark"]), obj))
            except (tk.TclError, KeyError):
                continue
        live_objects.sort(key=lambda x: tuple(map(int, x[0].split("."))))

        segments: list[dict] = []
        cursor = "1.0"
        for pos, obj in live_objects:
            chunk = self.editor.get(cursor, pos)
            if chunk: segments.append({"kind": "text", "text": chunk})
            segments.append({"kind": obj["type"], **obj["getter"]()})
            cursor = pos
        tail = self.editor.get(cursor, "end-1c")
        if tail: segments.append({"kind": "text", "text": tail})
        return segments

    def write_file(self, path: str):
        if not HAS_DOCX:
            messagebox.showerror("Missing Dependency",
                "MathsNotes needs python-docx to save .docx files.\n\nInstall with:\n  pip install python-docx")
            return
        try:
            segments = self.get_document_segments()
            doc = Document()
            normal_style = doc.styles["Normal"]
            normal_style.font.name = self.config.get("font_family", "Times New Roman")
            normal_style.font.size = Pt(self.config.get("font_size", 14))

            sidecar_objects: list[dict] = []
            body_index = 0  # track body child index manually because python-docx doesn't expose this directly

            for seg in segments:
                kind = seg["kind"]
                if kind == "text":
                    for line in seg["text"].split("\n"):
                        doc.add_paragraph(line)
                        body_index += 1
                elif kind == "fraction":
                    n = seg.get("numerator","")
                    d = seg.get("denominator","")
                    doc.add_paragraph(f"({n})/({d})") # even when the formatting isn't shown like in the app, it shows a plaintext version that still explains 
                                                        # what's going on (e.g. if you opened in word) like (x)/(y)
                    sidecar_objects.append({"type":"fraction","body_index":body_index,"numerator":n,"denominator":d})
                    body_index += 1
                elif kind == "table":
                    table_data = seg.get("data", [])
                    has_header = seg.get("has_header", False)
                    if not table_data: continue
                    num_cols = max(len(r) for r in table_data)
                    word_table = doc.add_table(rows=len(table_data), cols=num_cols)
                    word_table.style = "Table Grid"
                    for r, row in enumerate(table_data):
                        for c in range(num_cols):
                            cell = word_table.rows[r].cells[c]
                            cell.text = row[c] if c < len(row) else ""
                            if has_header and r == 0: # bold the header row
                                for para in cell.paragraphs:
                                    for run in para.runs: run.bold = True
                    sidecar_objects.append({"type":"table","body_index":body_index,"has_header":has_header,"data":table_data})
                    body_index += 1

            tmp_path = path + ".tmp"
            doc.save(tmp_path)
            self.inject_sidecar(tmp_path, {"version": self.SIDECAR_VER, "objects": sidecar_objects})
            os.replace(tmp_path, path) # atomic replace so a crash mid-write doesn't corrupt the file

            self.last_saved = time.strftime("%H:%M:%S")
            self.root.title(f"MathsNotes — {os.path.basename(path)}")
            self.root.after(0, lambda: self.save_status_label.configure(text=f"Saved at {self.last_saved}"))
        except Exception as e:
            messagebox.showerror("Save Error", f"Could not save file:\n{e}")

    def inject_sidecar(self, docx_path: str, payload: dict):
        # rewrite the docx zip to include the sidecar, replacing the old copy if it exists
        tmp_path = docx_path + ".sc.tmp"
        data = json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")
        with zipfile.ZipFile(docx_path, "r") as zip_in:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zip_out:
                for item in zip_in.infolist():
                    if item.filename == self.SIDECAR_PART: continue # skip old sidecar
                    zip_out.writestr(item, zip_in.read(item.filename))
                zip_out.writestr(self.SIDECAR_PART, data) # write the new sidecar
        os.replace(tmp_path, docx_path)
# autosave

    def autosave_loop(self):
        if not self.config.get("autosave_enabled", True) or self.file_path is None:
            return
        interval = self.config.get("autosave_interval", 30)
        with self.save_lock: self.write_file(self.file_path)
        self.autosave_timer = threading.Timer(interval, self.autosave_loop) # reschedule itself
        self.autosave_timer.daemon = True
        self.autosave_timer.start()

    def start_autosave(self):
        self.stop_autosave() # cancel any running timer before starting a fresh one
        if not self.config.get("autosave_enabled", True): return
        self.autosave_timer = threading.Timer(self.config.get("autosave_interval", 30), self.autosave_loop)
        self.autosave_timer.daemon = True
        self.autosave_timer.start()

    def stop_autosave(self):
        if self.autosave_timer:
            self.autosave_timer.cancel()
            self.autosave_timer = None
# table

    def open_table_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Insert Table — Step 1: Dimensions")
        dlg.transient(self.root); dlg.grab_set(); dlg.resizable(False, False)
        tk.Label(dlg, text="Rows:").grid(   row=0, column=0, padx=12, pady=8, sticky="e")
        rows_var = tk.IntVar(value=3)
        tk.Spinbox(dlg, from_=1, to=50, width=5, textvariable=rows_var).grid(row=0, column=1, padx=8, sticky="w")
        tk.Label(dlg, text="Columns:").grid(row=1, column=0, padx=12, pady=8, sticky="e")
        cols_var = tk.IntVar(value=3)
        tk.Spinbox(dlg, from_=1, to=12, width=5, textvariable=cols_var).grid(row=1, column=1, padx=8, sticky="w")
        header_var = tk.BooleanVar(value=True)
        tk.Checkbutton(dlg, text="First row is header", variable=header_var).grid(
            row=2, column=0, columnspan=2, padx=12, pady=(0,6), sticky="w")
        def go_to_step_two():
            r = rows_var.get()
            c = cols_var.get()
            h = header_var.get()
            dlg.destroy()
            self.open_table_fill_dialog(r, c, h)
        tk.Button(dlg, text="Next →", width=12, command=go_to_step_two).grid(row=3, column=0, columnspan=2, pady=12)

    def open_table_fill_dialog(self, rows: int, cols: int, has_header: bool):
        dlg = tk.Toplevel(self.root)
        dlg.title("Insert Table — Step 2: Cell Contents")
        dlg.transient(self.root); dlg.grab_set()
        tk.Label(dlg,
                 text="Fill in cell contents below. The inserted table will be a real "
                      "gridded widget — every cell remains editable after insertion.",
                 wraplength=520, justify="left", font=("Arial", 9)
                 ).pack(padx=10, pady=(10,6), anchor="w")
        grid_frame = tk.Frame(dlg); grid_frame.pack(padx=10, pady=6)
        cell_entries: list[list[tk.Entry]] = []
        for r in range(rows):
            row_entries = []
            for c in range(cols):
                entry = tk.Entry(grid_frame, width=14,
                             font=("Arial", 10, "bold" if (has_header and r==0) else "normal"))
                entry.grid(row=r, column=c, padx=2, pady=2)
                if has_header and r == 0: entry.insert(0, f"Col {c+1}") # pre-fill header cells
                row_entries.append(entry)
            cell_entries.append(row_entries)
        def do_insert():
            data = [[entry.get() for entry in row] for row in cell_entries]
            dlg.destroy()
            self.make_table_widget(data, has_header)
        button_frame = tk.Frame(dlg); button_frame.pack(pady=10)
        tk.Button(button_frame, text="← Back", width=10,
                  command=lambda: (dlg.destroy(), self.open_table_dialog())).pack(side="left", padx=4)
        tk.Button(button_frame, text="Insert Table", width=14, command=do_insert).pack(side="left", padx=4)

    def make_table_widget(self, data: list[list[str]], has_header: bool):
        if not data: return
        num_cols = max(len(r) for r in data)
        for row in data:
            while len(row) < num_cols: row.append("") # normalise so all rows have equal column count

        t = self.theme
        header_bg = "#EEF2F7" if self.config.get("theme","light") == "light" else "#27293D"
        table_frame = tk.Frame(self.editor, background=t["border"], bd=0,
                          highlightthickness=1, highlightbackground=t["border"])
        if not hasattr(self, "table_frames"):
            self.table_frames: list[tk.Frame] = []
        self.table_frames.append(table_frame)

        body_font   = (self.config.get("font_family","Times New Roman"), self.config.get("font_size",14))
        header_font = body_font + ("bold",)
        entry_grid: list[list[tk.Entry]] = []
        for r, row in enumerate(data):
            entry_row: list[tk.Entry] = []
            for c, val in enumerate(row):
                is_header = has_header and r == 0
                cell_entry = tk.Entry(table_frame,
                               font=header_font if is_header else body_font,
                               relief="flat", bd=0, highlightthickness=0,
                               background=header_bg if is_header else t["text_bg"],
                               foreground=t["text_fg"], insertbackground=t["fg"],
                               justify="left")
                cell_entry.insert(0, val)
                cell_entry.configure(width=max(8, min(24, len(val)+2)))
                cell_entry.grid(row=r, column=c, padx=(0,1), pady=(0,1), ipadx=4, ipady=2, sticky="nsew")
                entry_row.append(cell_entry)
            entry_grid.append(entry_row)
        for c in range(num_cols):
            table_frame.grid_columnconfigure(c, weight=1, uniform="tablecol") # equal column widths

        self.editor.window_create("insert", window=table_frame, align="top")
        self.editor.insert("insert", "\n") # newline after so subsequent text starts on the next line

        mark_name = f"obj_{self.object_counter}"
        self.object_counter += 1
        self.editor.mark_set(mark_name, "insert - 2 chars")
        self.editor.mark_gravity(mark_name, "left")
        self.embedded_objects.append({
            "type": "table", "mark": mark_name, "widget": table_frame, "has_header": has_header,
            "getter": lambda g=entry_grid, h=has_header: {
                "has_header": h,
                "data": [[cell.get() for cell in row] for row in g],
            },
        })
# chart

    def open_chart_dialog(self):
        if not HAS_MPL:
            messagebox.showwarning("Matplotlib Missing",
                                   "Install matplotlib to use chart insertion:\n  pip install matplotlib")
            return
        dlg = tk.Toplevel(self.root)
        dlg.title("Insert Chart"); dlg.transient(self.root)
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text="Chart Type:").grid(row=0, column=0, padx=12, pady=8, sticky="e")
        chart_type_var = tk.StringVar(value="Line")
        ttk.Combobox(dlg, textvariable=chart_type_var, values=["Line","Bar","Scatter","Pie"],
                     state="readonly", width=10).grid(row=0, column=1, padx=8)
        tk.Label(dlg, text="X values (comma-separated):").grid(row=1, column=0, padx=12, pady=4, sticky="e") # basic form controls for ideal formatting
        x_entry = tk.Entry(dlg, width=28); x_entry.insert(0, "1, 2, 3, 4, 5"); x_entry.grid(row=1, column=1, padx=8)
        tk.Label(dlg, text="Y values (comma-separated):").grid(row=2, column=0, padx=12, pady=4, sticky="e")
        y_entry = tk.Entry(dlg, width=28); y_entry.insert(0, "2, 4, 6, 8, 10"); y_entry.grid(row=2, column=1, padx=8)
        tk.Label(dlg, text="Title:").grid(row=3, column=0, padx=12, pady=4, sticky="e")
        title_entry = tk.Entry(dlg, width=28); title_entry.grid(row=3, column=1, padx=8)

        def do_insert():
            try:
                xs = [float(v.strip()) for v in x_entry.get().split(",") if v.strip()]
                ys = [float(v.strip()) for v in y_entry.get().split(",") if v.strip()]
            except ValueError:
                messagebox.showerror("Input Error", "Please enter valid numbers.", parent=dlg)
                return
            chart_type = chart_type_var.get()
            title = title_entry.get() or "Chart"
            dlg.destroy()
            fig, ax = plt.subplots(figsize=(5,3), dpi=100)
            if   chart_type == "Line":    ax.plot(xs, ys, marker="o")
            elif chart_type == "Bar":     ax.bar(xs, ys)
            elif chart_type == "Scatter": ax.scatter(xs, ys)
            elif chart_type == "Pie":     ax.pie(ys, labels=[str(int(x)) for x in xs], autopct="%1.1f%%")
            ax.set_title(title); plt.tight_layout()
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            fig.savefig(tmp.name, format="png"); plt.close(fig); tmp.close()
            img = tk.PhotoImage(file=tmp.name)
            if not hasattr(self, "chart_images"): self.chart_images = []
            self.chart_images.append(img) # keep a reference so the image isn't garbage collected
            self.editor.image_create("insert", image=img)
            self.editor.insert("insert", "\n")

        tk.Button(dlg, text="Insert Chart", command=do_insert, width=14).grid(
            row=4, column=0, columnspan=2, pady=12)
# sym dlg

    def open_symbol_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Insert Symbol — Notation Selection")
        dlg.transient(self.root); dlg.grab_set(); dlg.geometry("520x540")
        tk.Label(dlg, text="Search and double-click a symbol to insert it:",
                 font=("Arial", 10)).pack(padx=12, pady=(10,4), anchor="w")
        search_var = tk.StringVar()
        search_entry = tk.Entry(dlg, textvariable=search_var)
        search_entry.pack(fill="x", padx=12, pady=(0,6))
        search_entry.focus_set()
        list_frame = tk.Frame(dlg); list_frame.pack(fill="both", expand=True, padx=8, pady=4)
        scrollbar = tk.Scrollbar(list_frame); scrollbar.pack(side="right", fill="y")
        symbol_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, font=("Courier New", 12), selectmode="browse")
        symbol_listbox.pack(fill="both", expand=True)
        scrollbar.config(command=symbol_listbox.yview)
        def refill(*_):
            symbol_listbox.delete(0, "end")
            query = search_var.get().lower().strip()
            for shortcut, symbol in self.charmap.items():
                if not query or query in shortcut.lower() or query in symbol.lower():
                    symbol_listbox.insert("end", f"  {symbol}    {shortcut}") # symbol first for visual scanning
        refill(); search_var.trace_add("write", refill)
        def do_insert(ev=None):
            sel = symbol_listbox.curselection()
            if sel:
                sym = symbol_listbox.get(sel[0]).strip().split()[0] # first token is the symbol
                self.editor.insert("insert", sym); dlg.destroy()
        symbol_listbox.bind("<Double-Button-1>", do_insert)
        search_entry.bind("<Return>", do_insert)
        button_frame = tk.Frame(dlg); button_frame.pack(pady=8)
        tk.Button(button_frame, text="Insert Selected", width=16, command=do_insert).pack(side="left", padx=4)
        tk.Button(button_frame, text="Close", width=10, command=dlg.destroy).pack(side="left", padx=4)

    def insert_from_combobox(self, ev=None):
        val = self.notation_var.get()
        if "→" in val:
            self.editor.insert("insert", val.split("→")[-1].strip()) # extract just the symbol
        self.notation_var.set("Insert symbol…")
        self.reload_notation_combobox()
        self.editor.focus_set()

    def insert_from_sidebar(self, ev=None):
        sel = self.sidebar_listbox.curselection()
        if sel:
            parts = self.sidebar_listbox.get(sel[0]).strip().split()
            if len(parts) >= 2 and parts[-1] in self.charmap.values():
                self.editor.insert("insert", parts[-1])
# settings

    def open_settings_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Settings"); dlg.transient(self.root)
        dlg.grab_set(); dlg.resizable(False, False)
        nb = ttk.Notebook(dlg); nb.pack(fill="both", expand=True, padx=8, pady=8)
        general_tab = tk.Frame(nb); nb.add(general_tab, text="General")
        autosave_enabled_var = tk.BooleanVar(value=self.config.get("autosave_enabled", True))
        tk.Checkbutton(general_tab, text="Enable Auto-Save", variable=autosave_enabled_var).grid(
            row=0, column=0, columnspan=2, sticky="w", padx=12, pady=(12,4))
        tk.Label(general_tab, text="Auto-Save Interval (seconds):").grid(row=1, column=0, sticky="e", padx=12, pady=4)
        autosave_interval_var = tk.IntVar(value=self.config.get("autosave_interval", 30))
        tk.Spinbox(general_tab, from_=10, to=600, width=6, textvariable=autosave_interval_var).grid(row=1, column=1, sticky="w", padx=8)
        tk.Label(general_tab, text="Default Font Size:").grid(row=2, column=0, sticky="e", padx=12, pady=4)
        font_size_settings_var = tk.IntVar(value=self.config.get("font_size", 14))
        tk.Spinbox(general_tab, from_=8, to=72, width=6, textvariable=font_size_settings_var).grid(row=2, column=1, sticky="w", padx=8)
        tk.Label(general_tab, text="Theme:").grid(row=3, column=0, sticky="e", padx=12, pady=4)
        theme_var = tk.StringVar(value=self.config.get("theme", "light"))
        ttk.Combobox(general_tab, textvariable=theme_var, values=["light","dark"],
                     state="readonly", width=10).grid(row=3, column=1, sticky="w", padx=8)
        def commit_settings():
            self.config["autosave_enabled"]  = autosave_enabled_var.get()
            self.config["autosave_interval"] = autosave_interval_var.get()
            self.config["font_size"]         = font_size_settings_var.get()
            self.config["theme"]             = theme_var.get()
            self.save_config(); self.set_theme(theme_var.get())
            self.stop_autosave()
            if self.file_path: self.start_autosave()
            self.font_size_var.set(font_size_settings_var.get()); self.apply_font_globally()
            dlg.destroy()
        tk.Button(dlg, text="Save Settings", command=commit_settings, width=14).pack(pady=10)
# charmap ed

    def open_charmap_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("Edit Character Map"); dlg.transient(self.root)
        dlg.grab_set(); dlg.geometry("520x460")
        tk.Label(dlg,
                 text="Double-click a row to edit. Changes saved immediately to charmap.json.",
                 font=("Arial", 9), wraplength=480, justify="left"
                 ).pack(padx=12, pady=(8,4), anchor="w")
        tree_frame = tk.Frame(dlg); tree_frame.pack(fill="both", expand=True, padx=8)
        tree = ttk.Treeview(tree_frame, columns=("shortcut","symbol"), show="headings", height=14)
        tree.heading("shortcut", text="Shortcut"); tree.heading("symbol", text="Symbol")
        tree.column("shortcut", width=200);        tree.column("symbol",  width=260)
        tree_scroll = tk.Scrollbar(tree_frame, command=tree.yview); tree.configure(yscrollcommand=tree_scroll.set)
        tree.pack(side="left", fill="both", expand=True); tree_scroll.pack(side="right", fill="y")
        for shortcut, symbol in self.charmap.items():
            tree.insert("", "end", values=(shortcut, symbol))

        def on_double_click(ev):
            item = tree.identify_row(ev.y)
            if not item: return
            old_shortcut, old_symbol = tree.item(item, "values")
            new_shortcut = simpledialog.askstring("Edit Shortcut", f"Shortcut (was '{old_shortcut}'):",
                                        initialvalue=old_shortcut, parent=dlg)
            if new_shortcut is None: return
            new_symbol = simpledialog.askstring("Edit Symbol", f"Symbol (was '{old_symbol}'):",
                                        initialvalue=old_symbol, parent=dlg)
            if new_symbol is None: return
            del self.charmap[old_shortcut]; self.charmap[new_shortcut] = new_symbol
            self.max_shortcut_len = max((len(k) for k in self.charmap), default=0)
            self.save_charmap(); tree.item(item, values=(new_shortcut, new_symbol))
            self.refill_sidebar(); self.reload_notation_combobox()

        def add_entry():
            new_shortcut = simpledialog.askstring("New Shortcut", "Shortcut text:", parent=dlg)
            if not new_shortcut: return
            new_symbol = simpledialog.askstring("New Symbol", "Unicode symbol:", parent=dlg)
            if not new_symbol: return
            self.charmap[new_shortcut] = new_symbol
            self.max_shortcut_len = max((len(k) for k in self.charmap), default=0)
            self.save_charmap(); tree.insert("", "end", values=(new_shortcut, new_symbol))
            self.refill_sidebar(); self.reload_notation_combobox()

        def del_entry():
            sel = tree.selection()
            if not sel: return
            shortcut = tree.item(sel[0], "values")[0]
            if messagebox.askyesno("Delete", f"Delete shortcut '{shortcut}'?", parent=dlg):
                del self.charmap[shortcut]; self.save_charmap()
                tree.delete(sel[0]); self.refill_sidebar(); self.reload_notation_combobox()

        def def_reset():
            if messagebox.askyesno("Reset", "Replace the current map with the built-in defaults?", parent=dlg):
                self.charmap = dict(defaultcm)
                self.max_shortcut_len = max((len(k) for k in self.charmap), default=0)
                self.save_charmap()
                for item in tree.get_children(): tree.delete(item)
                for shortcut, symbol in self.charmap.items(): tree.insert("", "end", values=(shortcut, symbol))
                self.refill_sidebar(); self.reload_notation_combobox()

        tree.bind("<Double-Button-1>", on_double_click)
        button_frame = tk.Frame(dlg); button_frame.pack(pady=6)
        tk.Button(button_frame, text="Add",           width=8,  command=add_entry).pack(   side="left", padx=4)
        tk.Button(button_frame, text="Delete",        width=8,  command=del_entry).pack(side="left", padx=4)
        tk.Button(button_frame, text="Reset Defaults",width=14, command=def_reset).pack(side="left", padx=4)
# sidebar tog

    def sidebar_on(self):
        if self.sidebar_visible:
            self.paned.forget(self.sidebar_frame)
            self.sidebar_visible = False
        else:
            self.paned.add(self.sidebar_frame, stretch="never", minsize=220)
            self.sidebar_visible = True
# help

    def show_help(self):
        win = tk.Toplevel(self.root)
        win.title("MathsNotes — Help Manual"); win.geometry("680x560")
        txt = scrolledtext.ScrolledText(win, wrap="word", font=("Arial", 11),
                                        relief="flat", padx=16, pady=12)
        txt.pack(fill="both", expand=True)
        manual = """MathsNotes — Help Manual
═══════════════════════════════════════════════════

MathsNotes is a digital maths notation editor.
Type your notes normally. Some key sequences turn into maths symbols automatically. You can see the shortcuts in the sidebar.

GETTING STARTED
───────────────
1. Open the app and start typing.
2. Use the Character Map panel on the right to see all available shortcuts.
3. Type a shortcut (e.g. "pi") and it will instantly become your desired
character, e.g. "π".
4. Save your file with Ctrl+S. After the first save, auto-save activates.

CHARACTER MAP — basic substitutions
────────────────────────────────────
+-       → ±        *        → ×        -:-      → ÷
=/=      → ≠        >_       → ≥        <_       → ≤
->       → ⇒        <->      → ⇔        ~~       → ≈
sqrt     → √        inf      → ∞        integral → ∫
sum      → ∑        partial  → ∂        nabla    → ∇
pi       → π        theta    → θ        omega    → ω
alpha    → α        beta     → β        sigma    → σ
Re       → ℝ        Nat      → ℕ        Com      → ℂ
degrees  → °        forall   → ∀        exists   → ∃

(See the full list and search in the sidebar.)

MODAL TRIGGERS — special interactive shortcuts
────────────────────────────────────────────────
/  Fraction mode. After a digit (or a bracket like "(a+b)"), 
   typing '/' immediately turns the preceding string
   into the numerator of a real two-line stacked fraction
   and offers an empty denominator slot for you to type
   into. Press Enter, Tab, Escape or the right arrow to
   return to the main editor. Both halves of the fraction 
   stay editable.

^  Superscript mode. Press ^, then type 0-9, +, -, n, i…
   Each character is converted to its Unicode superscript.
   This would most commonly be used for powers (squaring,
   cubing…) Any non-mappable character exits the mode.

_  Subscript mode. Same as ^ but for subscripts. Press 
   "_" to activate.
   Note: ">_" and "<_" still produce ≥ and ≤.

INSERT MENU
────────────
• Table     — Two-step generator: pick dimensions, fill cells,
              insert an editable table.
• Chart     — Plot data according to the type of chart you want, it inserts the chart at the cursor.
• Symbol    — Searchable picker for every entry in the character map.
• Fraction  — Insert a fraction without typing '/'.

NOTATION SELECTION (toolbar)
─────────────────────────────
The "Notation:" dropdown in the toolbar lists every symbol
in the active character map. Type to filter, then Enter or
click to insert at the cursor.

KEYBOARD SHORTCUTS
───────────────────
Ctrl+N   New file        Ctrl+Z   Undo
Ctrl+O   Open file       Ctrl+Y   Redo
Ctrl+S   Save            Ctrl+B   Toggle sidebar
Ctrl+Shift+S  Save As    F1       This help window

SETTINGS
─────────
• Auto-save interval (10–600 seconds) and on/off toggle
• Default font size and family
• Light or Dark theme
• Settings > Edit Character Map: add, edit, delete, or
  reset shortcuts. Stored in charmap.json.
"""
        txt.insert("1.0", manual)
        txt.configure(state="disabled")

    def show_abt(self):
        messagebox.showinfo(
            "About MathsNotes",
            "MathsNotes — A Digital Mathematics Notation Editor\n\n"
            "A-Level Computer Science NEA Project\n"
            "Python 3 · tkinter · threading · matplotlib\n\n"
            "Libraries: tkinter, threading, re, json, os, matplotlib, python-docx"
        )
# utils

    def update_status(self, msg=""):
        line, col = self.editor.index("insert").split(".")
        self.status_label.configure(text=f"Ln {line}  Col {int(col)+1}   {msg}")

    def check_unsaved(self) -> bool:
        if not self.editor.get("1.0", "end-1c").strip():
            return True # empty document, nothing to lose
        ans = messagebox.askyesnocancel("Unsaved Changes",
                                        "You have unsaved changes. Save before continuing?")
        if ans is None: return False # user hit Cancel, abort the operation
        if ans: self.save_file()
        return True

    def on_close(self):
        self.stop_autosave() # cancel timer before destroying so the thread doesn't try to write after teardown
        if not self.check_unsaved(): return
        self.root.destroy() # die

    def _read_json(self, path: str, fallback: dict) -> dict:
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) # converts json file to dictionary for reference and application
            return data if isinstance(data, dict) else fallback
        except Exception:
            return fallback # don't crash if the file is missing or malformed

    def save_config(self):
        try:
            with open(configjs, "w", encoding="utf-8") as f:
                json.dump(self.config, f, indent=4) # saves from memory back into json
        except Exception:
            pass # don't crash at all costs

    def save_charmap(self):
        try:
            with open(charmapjs, "w", encoding="utf-8") as f:
                json.dump(self.charmap, f, indent=4, ensure_ascii=False) # ensure_ascii=False keeps unicode symbols readable
        except Exception:
            pass # don't crash at all costs
# equation eval
# only handles very simple numeric expressions

def try_evaluate_equation(line: str):
    line = line.strip().rstrip("=").strip() # strip trailing = so "3+4=" works as well as "3+4"
    if not line:
        return None
    safe = re.sub(r"[\d\s\+\-\*\/\.\(\)\^]", "", line)
    if safe:
        return None # reject anything with non-numeric characters to avoid arbitrary code execution
    try:
        result = eval(line.replace("^","**"), {"__builtins__": {}}, {}) # no builtins = sandboxed
        if isinstance(result, (int, float)) and not math.isnan(result) and not math.isinf(result):
            return f"= {result}"
    except Exception:
        pass
    return None


def main():
    root = tk.Tk()
    root.title("MathsNotes")
    try:
        root.iconbitmap(default="")
    except Exception:
        pass
    MathsNotes(root)
    root.mainloop() # launchpad


if __name__ == "__main__":
    main() # activate launchpad
